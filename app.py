"""FastAPI web application for Research Agent System."""

import os
import uuid
from datetime import datetime, timedelta
from typing import Optional
from pathlib import Path

from fastapi import FastAPI, Request, Form, HTTPException, Depends, status
from fastapi.responses import HTMLResponse, JSONResponse, FileResponse, RedirectResponse
from fastapi.templating import Jinja2Templates
from fastapi.staticfiles import StaticFiles
from starlette.middleware.sessions import SessionMiddleware
from sqlalchemy.orm import Session
from dotenv import load_dotenv
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from authlib.integrations.starlette_client import OAuth, OAuthError

from agents.orchestrator_agent import OrchestratorAgent
from database import init_db, get_db, User, ResearchHistory, engine
from auth import (
    verify_password, get_password_hash, create_access_token,
    get_current_user, get_current_user_optional, SECRET_KEY, ALGORITHM
)
from jose import jwt, JWTError
from sqladmin import Admin, ModelView
from sqladmin.authentication import AuthenticationBackend

# Load environment variables
load_dotenv()

# Initialize database
init_db()

# Initialize FastAPI app
app = FastAPI(title="Research Agent System", version="1.0.0")

# Add session middleware for SQLAdmin authentication (must be added before creating Admin)
# Use a consistent secret key for sessions - must be the same for both middleware and AdminAuth
SESSION_SECRET_KEY = SECRET_KEY if SECRET_KEY != "your-secret-key-change-this-in-production" else "admin-session-secret-key-change-in-production"
app.add_middleware(SessionMiddleware, secret_key=SESSION_SECRET_KEY, max_age=86400)  # 24 hours

# SQLAdmin setup for web-based database management
class AdminAuth(AuthenticationBackend):
    """Simple authentication for admin panel."""
    
    async def login(self, request: Request) -> bool:
        """Handle login form submission."""
        form = await request.form()
        username = form.get("username")
        password = form.get("password")
        
        # Simple admin authentication (change these in production!)
        admin_username = os.getenv("ADMIN_USERNAME", "admin")
        admin_password = os.getenv("ADMIN_PASSWORD", "admin")
        
        if username == admin_username and password == admin_password:
            # Store authentication in session
            request.session["authenticated"] = True
            request.session["username"] = username
            # Debug: Print session state
            print(f"[ADMIN LOGIN] Session after login: {dict(request.session)}")
            return True
        return False

    async def logout(self, request: Request) -> bool:
        """Handle logout."""
        request.session.clear()
        return True

    async def authenticate(self, request: Request) -> bool:
        """Check if user is authenticated."""
        # Check if user is authenticated in session
        # SQLAdmin calls this method on every request to protected routes
        authenticated = request.session.get("authenticated", False)
        # Debug: Print authentication check
        print(f"[ADMIN AUTH] Checking auth - authenticated: {authenticated}, session keys: {list(request.session.keys())}")
        return bool(authenticated)


# Create admin interface with authentication
# The secret_key must match the SessionMiddleware secret_key for sessions to work
authentication_backend = AdminAuth(secret_key=SESSION_SECRET_KEY)
admin = Admin(app, engine, authentication_backend=authentication_backend)


# Admin views for database models
class UserAdmin(ModelView, model=User):
    """Admin interface for User model."""
    column_list = [User.id, User.email, User.username, User.full_name, User.provider, User.created_at, User.is_active]
    column_searchable_list = [User.email, User.username]
    column_sortable_list = [User.id, User.email, User.created_at]
    can_create = True
    can_edit = True
    can_delete = True
    can_view_details = True


class ResearchHistoryAdmin(ModelView, model=ResearchHistory):
    """Admin interface for ResearchHistory model."""
    column_list = [ResearchHistory.id, ResearchHistory.user_id, ResearchHistory.topic, ResearchHistory.created_at, ResearchHistory.sources_count, ResearchHistory.queries_count]
    column_searchable_list = [ResearchHistory.topic]
    column_sortable_list = [ResearchHistory.id, ResearchHistory.created_at]
    can_create = False  # Research history is created by the system
    can_edit = True
    can_delete = True
    can_view_details = True
    # Show full content in details view
    column_details_list = [ResearchHistory.id, ResearchHistory.user_id, ResearchHistory.topic, 
                          ResearchHistory.requirements, ResearchHistory.draft, ResearchHistory.improved_draft,
                          ResearchHistory.changes_summary, ResearchHistory.sources_count, ResearchHistory.queries_count,
                          ResearchHistory.md_filename, ResearchHistory.docx_filename, ResearchHistory.created_at]


# Register admin views
admin.add_view(UserAdmin)
admin.add_view(ResearchHistoryAdmin)

# Setup templates and static files
BASE_DIR = Path(__file__).parent
templates = Jinja2Templates(directory=str(BASE_DIR / "templates"))
static_dir = BASE_DIR / "static"
static_dir.mkdir(exist_ok=True)
app.mount("/static", StaticFiles(directory=str(static_dir)), name="static")

# Create downloads directory
downloads_dir = BASE_DIR / "downloads"
downloads_dir.mkdir(exist_ok=True)

# OAuth setup
oauth = OAuth()
GOOGLE_CLIENT_ID = os.getenv("GOOGLE_CLIENT_ID", "")
GOOGLE_CLIENT_SECRET = os.getenv("GOOGLE_CLIENT_SECRET", "")
FACEBOOK_CLIENT_ID = os.getenv("FACEBOOK_CLIENT_ID", "")
FACEBOOK_CLIENT_SECRET = os.getenv("FACEBOOK_CLIENT_SECRET", "")

if GOOGLE_CLIENT_ID and GOOGLE_CLIENT_SECRET:
    oauth.register(
        name='google',
        client_id=GOOGLE_CLIENT_ID,
        client_secret=GOOGLE_CLIENT_SECRET,
        server_metadata_url='https://accounts.google.com/.well-known/openid-configuration',
        client_kwargs={'scope': 'openid email profile'}
    )

if FACEBOOK_CLIENT_ID and FACEBOOK_CLIENT_SECRET:
    oauth.register(
        name='facebook',
        client_id=FACEBOOK_CLIENT_ID,
        client_secret=FACEBOOK_CLIENT_SECRET,
        server_metadata_url='https://www.facebook.com/.well-known/openid-configuration',
        client_kwargs={'scope': 'email public_profile'}
    )


@app.get("/", response_class=HTMLResponse)
async def home(request: Request, db: Session = Depends(get_db)):
    """Render the home page."""
    current_user = await get_current_user_optional(request, db)
    return templates.TemplateResponse("index.html", {
        "request": request,
        "user": current_user,
        "google_enabled": bool(GOOGLE_CLIENT_ID),
        "facebook_enabled": bool(FACEBOOK_CLIENT_ID)
    })


@app.get("/login", response_class=HTMLResponse)
async def login_page(request: Request):
    """Render the login page."""
    return templates.TemplateResponse("login.html", {
        "request": request,
        "google_enabled": bool(GOOGLE_CLIENT_ID),
        "facebook_enabled": bool(FACEBOOK_CLIENT_ID)
    })


@app.get("/register", response_class=HTMLResponse)
async def register_page(request: Request):
    """Render the registration page."""
    return templates.TemplateResponse("register.html", {
        "request": request,
        "google_enabled": bool(GOOGLE_CLIENT_ID),
        "facebook_enabled": bool(FACEBOOK_CLIENT_ID)
    })


@app.post("/api/auth/register")
async def register(
    email: str = Form(...),
    username: str = Form(...),
    password: str = Form(...),
    full_name: Optional[str] = Form(None),
    db: Session = Depends(get_db)
):
    """Register a new user."""
    try:
        # Validate inputs
        email = email.strip().lower() if email else ""
        username = username.strip() if username else ""
        
        if not email:
            raise HTTPException(status_code=400, detail="Email is required")
        if not username:
            raise HTTPException(status_code=400, detail="Username is required")
        if not password:
            raise HTTPException(status_code=400, detail="Password is required")
        
        # Validate password length
        if len(password) < 6:
            raise HTTPException(status_code=400, detail="Password must be at least 6 characters")
        if len(password.encode('utf-8')) > 72:
            raise HTTPException(status_code=400, detail="Password is too long (maximum 72 bytes)")
        
        # Check if user already exists
        existing_user = db.query(User).filter(User.email == email).first()
        if existing_user:
            raise HTTPException(status_code=400, detail="Email already registered")
        
        existing_username = db.query(User).filter(User.username == username).first()
        if existing_username:
            raise HTTPException(status_code=400, detail="Username already taken")
        
        # Create new user
        hashed_password = get_password_hash(password)
        user = User(
            email=email,
            username=username,
            hashed_password=hashed_password,
            full_name=full_name.strip() if full_name else None,
            provider="local"
        )
        db.add(user)
        db.commit()
        db.refresh(user)
        
        # Create access token (sub must be a string)
        access_token = create_access_token(data={"sub": str(user.id)})
        
        return JSONResponse({
            "success": True,
            "access_token": access_token,
            "token_type": "bearer",
            "user": {
                "id": user.id,
                "email": user.email,
                "username": user.username,
                "full_name": user.full_name
            }
        })
    except HTTPException:
        raise
    except Exception as e:
        import traceback
        error_detail = f"Registration failed: {str(e)}"
        print(f"Registration error: {error_detail}")
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=error_detail)


@app.post("/api/auth/login")
async def login(
    username: str = Form(...),
    password: str = Form(...),
    db: Session = Depends(get_db)
):
    """Login with email/username and password."""
    # Try email first, then username
    user = db.query(User).filter(
        (User.email == username) | (User.username == username)
    ).first()
    
    if not user or not user.hashed_password:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Incorrect email/username or password",
            headers={"WWW-Authenticate": "Bearer"},
        )
    
    if not verify_password(password, user.hashed_password):
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Incorrect email/username or password",
            headers={"WWW-Authenticate": "Bearer"},
        )
    
    access_token = create_access_token(data={"sub": str(user.id)})
    
    # Return JSONResponse to ensure proper formatting
    return JSONResponse({
        "access_token": access_token,
        "token_type": "bearer",
        "user": {
            "id": user.id,
            "email": user.email,
            "username": user.username,
            "full_name": user.full_name
        }
    })


@app.get("/api/auth/google")
async def google_login(request: Request):
    """Initiate Google OAuth login."""
    if not GOOGLE_CLIENT_ID:
        raise HTTPException(status_code=400, detail="Google OAuth not configured")
    
    redirect_uri = request.url_for('google_callback')
    return await oauth.google.authorize_redirect(request, redirect_uri)


@app.get("/api/auth/google/callback")
async def google_callback(request: Request, db: Session = Depends(get_db)):
    """Handle Google OAuth callback."""
    try:
        token = await oauth.google.authorize_access_token(request)
        user_info = token.get('userinfo')
        
        if not user_info:
            raise HTTPException(status_code=400, detail="Failed to get user info")
        
        email = user_info.get('email')
        provider_id = user_info.get('sub')
        full_name = user_info.get('name')
        
        # Find or create user
        user = db.query(User).filter(User.email == email).first()
        if not user:
            user = User(
                email=email,
                username=email.split('@')[0],
                full_name=full_name,
                provider="google",
                provider_id=provider_id
            )
            db.add(user)
            db.commit()
            db.refresh(user)
        elif user.provider != "google":
            raise HTTPException(status_code=400, detail="Email already registered with different provider")
        
        access_token = create_access_token(data={"sub": str(user.id)})
        
        # Redirect to home with token
        response = RedirectResponse(url="/?token=" + access_token)
        return response
        
    except OAuthError as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.get("/api/auth/facebook")
async def facebook_login(request: Request):
    """Initiate Facebook OAuth login."""
    if not FACEBOOK_CLIENT_ID:
        raise HTTPException(status_code=400, detail="Facebook OAuth not configured")
    
    redirect_uri = request.url_for('facebook_callback')
    return await oauth.facebook.authorize_redirect(request, redirect_uri)


@app.get("/api/auth/facebook/callback")
async def facebook_callback(request: Request, db: Session = Depends(get_db)):
    """Handle Facebook OAuth callback."""
    try:
        token = await oauth.facebook.authorize_access_token(request)
        user_info = token.get('userinfo')
        
        if not user_info:
            raise HTTPException(status_code=400, detail="Failed to get user info")
        
        email = user_info.get('email')
        provider_id = user_info.get('sub')
        full_name = user_info.get('name')
        
        # Find or create user
        user = db.query(User).filter(User.email == email).first()
        if not user:
            user = User(
                email=email,
                username=email.split('@')[0] if email else f"fb_{provider_id}",
                full_name=full_name,
                provider="facebook",
                provider_id=provider_id
            )
            db.add(user)
            db.commit()
            db.refresh(user)
        elif user.provider != "facebook":
            raise HTTPException(status_code=400, detail="Email already registered with different provider")
        
        access_token = create_access_token(data={"sub": str(user.id)})
        
        # Redirect to home with token
        response = RedirectResponse(url="/?token=" + access_token)
        return response
        
    except OAuthError as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.get("/api/auth/me")
async def get_current_user_info(request: Request, db: Session = Depends(get_db)):
    """Get current user information."""
    # Try to get token from Authorization header or query parameter
    token = None
    auth_header = request.headers.get("Authorization")
    
    if auth_header and auth_header.startswith("Bearer "):
        token = auth_header.split(" ", 1)[1]  # Use split with maxsplit=1 to handle tokens with spaces
    else:
        token = request.query_params.get("token")
    
    if not token:
        raise HTTPException(
            status_code=401, 
            detail="Not authenticated",
            headers={"WWW-Authenticate": "Bearer"}
        )
    
    try:
        # Decode and verify token
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        user_id_str = payload.get("sub")
        
        if user_id_str is None:
            raise HTTPException(status_code=401, detail="Invalid token: missing user ID")
        
        # Convert string to int for database query
        try:
            user_id = int(user_id_str)
        except (ValueError, TypeError):
            raise HTTPException(status_code=401, detail="Invalid token: user ID format error")
        
        user = db.query(User).filter(User.id == user_id).first()
        if user is None:
            raise HTTPException(status_code=401, detail="User not found")
        if not user.is_active:
            raise HTTPException(status_code=401, detail="User is inactive")
        return {
            "id": user.id,
            "email": user.email,
            "username": user.username,
            "full_name": user.full_name,
            "provider": user.provider
        }
    except JWTError as e:
        raise HTTPException(status_code=401, detail=f"Token validation failed: {str(e)}")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=401, detail=f"Authentication error: {str(e)}")


@app.post("/api/auth/logout")
async def logout():
    """Logout (client-side token removal)."""
    return {"success": True, "message": "Logged out successfully"}


@app.post("/api/research")
async def submit_research(
    topic: str = Form(...),
    requirements: Optional[str] = Form(None),
    max_sources: Optional[int] = Form(None),
    max_query_length: Optional[int] = Form(200),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Submit a research request and return results."""
    try:
        # Get max_sources from env if not provided
        if max_sources is None:
            max_sources = int(os.getenv('TAVILY_MAX_SOURCES', '5'))
        
        # Initialize orchestrator
        orchestrator = OrchestratorAgent()
        
        # Execute research workflow
        results = orchestrator.execute_research_workflow(
            topic=topic,
            requirements=requirements,
            max_sources_per_query=max_sources,
            max_query_length=max_query_length
        )
        
        # Generate unique ID for this research
        research_id = str(uuid.uuid4())
        
        # Save results to downloads directory
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename_base = f"research_{timestamp}_{research_id[:8]}"
        
        # Save markdown
        md_path = downloads_dir / f"{filename_base}.md"
        with open(md_path, 'w', encoding='utf-8') as f:
            f.write(f"# Research Report\n\n")
            f.write(f"**Topic:** {topic}\n\n")
            if requirements:
                f.write(f"**Requirements:** {requirements}\n\n")
            f.write(f"**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
            f.write(f"**Queries Executed:** {results['queries_count']}\n")
            f.write(f"**Sources Collected:** {results['sources_count']}\n\n")
            f.write("---\n\n")
            f.write("## Initial Draft\n\n")
            f.write(results['draft'])
            f.write("\n\n---\n\n")
            f.write("## Improved Draft\n\n")
            f.write(results['improved_draft'])
            if results.get('changes_summary'):
                f.write("\n\n---\n\n")
                f.write("## Changes Summary\n\n")
                f.write(results['changes_summary'])
        
        # Save Word document
        docx_path = downloads_dir / f"{filename_base}.docx"
        create_word_document(
            docx_path,
            topic=topic,
            requirements=requirements,
            results=results,
            timestamp=datetime.now()
        )
        
        # Save to database
        research_history = ResearchHistory(
            user_id=current_user.id,
            topic=topic,
            requirements=requirements,
            draft=results['draft'],
            improved_draft=results['improved_draft'],
            changes_summary=results.get('changes_summary', ''),
            sources_count=results['sources_count'],
            queries_count=results['queries_count'],
            md_filename=f"{filename_base}.md",
            docx_filename=f"{filename_base}.docx"
        )
        db.add(research_history)
        db.commit()
        db.refresh(research_history)
        
        return JSONResponse({
            "success": True,
            "research_id": str(research_history.id),
            "topic": topic,
            "draft": results['draft'],
            "improved_draft": results['improved_draft'],
            "changes_summary": results.get('changes_summary', ''),
            "sources_count": results['sources_count'],
            "queries_count": results['queries_count'],
            "downloads": {
                "markdown": f"/api/download/md/{research_history.id}",
                "word": f"/api/download/docx/{research_history.id}"
            },
            "filename_base": filename_base
        })
        
    except Exception as e:
        return JSONResponse(
            {"success": False, "error": str(e)},
            status_code=500
        )


@app.get("/api/download/md/{research_id}")
async def download_markdown(
    research_id: int,
    request: Request,
    db: Session = Depends(get_db)
):
    """Download research results as Markdown."""
    current_user = await get_current_user_optional(request, db)
    if not current_user:
        raise HTTPException(status_code=401, detail="Authentication required")
    
    research = db.query(ResearchHistory).filter(
        ResearchHistory.id == research_id,
        ResearchHistory.user_id == current_user.id
    ).first()
    
    if not research:
        raise HTTPException(status_code=404, detail="Research not found")
    
    file_path = downloads_dir / research.md_filename
    
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found")
    
    return FileResponse(
        path=file_path,
        filename=research.md_filename,
        media_type="text/markdown"
    )


@app.get("/api/download/docx/{research_id}")
async def download_word(
    research_id: int,
    request: Request,
    db: Session = Depends(get_db)
):
    """Download research results as Word document."""
    current_user = await get_current_user_optional(request, db)
    if not current_user:
        raise HTTPException(status_code=401, detail="Authentication required")
    
    research = db.query(ResearchHistory).filter(
        ResearchHistory.id == research_id,
        ResearchHistory.user_id == current_user.id
    ).first()
    
    if not research:
        raise HTTPException(status_code=404, detail="Research not found")
    
    file_path = downloads_dir / research.docx_filename
    
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found")
    
    return FileResponse(
        path=file_path,
        filename=research.docx_filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


@app.get("/api/research/history")
async def get_research_history(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get user's research history."""
    history = db.query(ResearchHistory).filter(
        ResearchHistory.user_id == current_user.id
    ).order_by(ResearchHistory.created_at.desc()).all()
    
    return {
        "history": [
            {
                "id": r.id,
                "topic": r.topic,
                "created_at": r.created_at.isoformat(),
                "sources_count": r.sources_count,
                "queries_count": r.queries_count
            }
            for r in history
        ]
    }


def create_word_document(
    filepath: Path,
    topic: str,
    requirements: Optional[str],
    results: dict,
    timestamp: datetime
):
    """Create a Word document from research results."""
    doc = Document()
    
    # Title
    title = doc.add_heading('Research Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    doc.add_paragraph(f"Topic: {topic}")
    if requirements:
        doc.add_paragraph(f"Requirements: {requirements}")
    doc.add_paragraph(f"Generated: {timestamp.strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Queries Executed: {results['queries_count']}")
    doc.add_paragraph(f"Sources Collected: {results['sources_count']}")
    
    doc.add_paragraph()  # Blank line
    
    # Initial Draft
    doc.add_heading('Initial Draft', level=1)
    add_text_to_document(doc, results['draft'])
    
    doc.add_page_break()
    
    # Improved Draft
    doc.add_heading('Improved Draft', level=1)
    add_text_to_document(doc, results['improved_draft'])
    
    # Changes Summary
    if results.get('changes_summary'):
        doc.add_page_break()
        doc.add_heading('Changes Summary', level=1)
        add_text_to_document(doc, results['changes_summary'])
    
    doc.save(str(filepath))


def add_text_to_document(doc: Document, text: str):
    """Add formatted text to a Word document."""
    lines = text.split('\n')
    
    for line in lines:
        line = line.strip()
        
        if not line:
            doc.add_paragraph()
            continue
        
        # Check if it's a heading (starts with #)
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            heading_text = line.lstrip('#').strip()
            if heading_text:
                doc.add_heading(heading_text, level=min(level, 9))
        # Check if it's a list item
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line and line[0].isdigit() and '. ' in line[:5]:
            content = line.split('. ', 1)[1] if '. ' in line else line
            doc.add_paragraph(content, style='List Number')
        else:
            # Regular paragraph
            doc.add_paragraph(line)


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
